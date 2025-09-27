"""Local file loader for various document types."""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Set
from datetime import datetime

from ..utils.config import config
from ..utils.logging import get_logger

logger = get_logger(__name__)


class LocalFileLoader:
    """Loader for local files of various types."""
    
    def __init__(self):
        """Initialize local file loader."""
        self.supported_extensions = {
            '.txt': self._load_text_file,
            '.md': self._load_text_file,
            '.py': self._load_text_file,
            '.js': self._load_text_file,
            '.ts': self._load_text_file,
            '.html': self._load_text_file,
            '.css': self._load_text_file,
            '.json': self._load_text_file,
            '.yaml': self._load_text_file,
            '.yml': self._load_text_file,
            '.pdf': self._load_pdf_file,
            '.docx': self._load_docx_file,
        }
    
    def load_files(
        self,
        paths: Optional[List[str]] = None,
        file_types: Optional[List[str]] = None,
        exclude_patterns: Optional[List[str]] = None,
        recursive: bool = True
    ) -> List[Dict[str, Any]]:
        """Load files from specified paths.
        
        Args:
            paths: List of file/directory paths to scan
            file_types: List of file extensions to include (e.g., ['.pdf', '.txt'])
            exclude_patterns: List of patterns to exclude
            recursive: Whether to scan directories recursively
            
        Returns:
            List of document dictionaries with metadata
        """
        paths = paths or config.get('local_files.paths', ['~/Documents'])
        file_types = file_types or config.get('local_files.file_types', list(self.supported_extensions.keys()))
        exclude_patterns = exclude_patterns or config.get('local_files.exclude_patterns', [])
        
        # Expand user paths
        expanded_paths = [Path(path).expanduser() for path in paths]
        
        logger.info(f"Loading files from paths: {expanded_paths}")
        logger.info(f"File types: {file_types}")
        
        documents = []
        
        for base_path in expanded_paths:
            if not base_path.exists():
                logger.warning(f"Path does not exist: {base_path}")
                continue
            
            if base_path.is_file():
                # Single file
                doc = self._process_file(base_path, exclude_patterns)
                if doc:
                    documents.append(doc)
            else:
                # Directory
                file_paths = self._scan_directory(base_path, file_types, exclude_patterns, recursive)
                
                for file_path in file_paths:
                    try:
                        doc = self._process_file(file_path, exclude_patterns)
                        if doc:
                            documents.append(doc)
                    except Exception as e:
                        logger.warning(f"Failed to process file {file_path}: {e}")
                        continue
        
        logger.info(f"Successfully loaded {len(documents)} local files")
        return documents
    
    def _scan_directory(
        self,
        directory: Path,
        file_types: List[str],
        exclude_patterns: List[str],
        recursive: bool
    ) -> List[Path]:
        """Scan directory for files matching criteria.
        
        Args:
            directory: Directory to scan
            file_types: File extensions to include
            exclude_patterns: Patterns to exclude
            recursive: Whether to scan recursively
            
        Returns:
            List of file paths
        """
        file_paths = []
        
        try:
            if recursive:
                pattern = "**/*"
            else:
                pattern = "*"
            
            for file_path in directory.glob(pattern):
                if not file_path.is_file():
                    continue
                
                # Check file extension
                if file_path.suffix.lower() not in [ext.lower() for ext in file_types]:
                    continue
                
                # Check exclude patterns
                if self._should_exclude(file_path, exclude_patterns):
                    continue
                
                file_paths.append(file_path)
        
        except Exception as e:
            logger.error(f"Error scanning directory {directory}: {e}")
        
        return file_paths
    
    def _should_exclude(self, file_path: Path, exclude_patterns: List[str]) -> bool:
        """Check if file should be excluded based on patterns.
        
        Args:
            file_path: File path to check
            exclude_patterns: List of exclusion patterns
            
        Returns:
            True if file should be excluded
        """
        import fnmatch
        
        file_str = str(file_path)
        
        for pattern in exclude_patterns:
            if fnmatch.fnmatch(file_str, pattern):
                return True
            
            # Also check just the filename
            if fnmatch.fnmatch(file_path.name, pattern):
                return True
        
        return False
    
    def _process_file(self, file_path: Path, exclude_patterns: List[str]) -> Optional[Dict[str, Any]]:
        """Process a single file.
        
        Args:
            file_path: Path to the file
            exclude_patterns: Patterns to exclude
            
        Returns:
            Document data or None if failed/excluded
        """
        try:
            # Check if file should be excluded
            if self._should_exclude(file_path, exclude_patterns):
                return None
            
            # Get file stats
            stat = file_path.stat()
            
            # Get file extension and loader
            extension = file_path.suffix.lower()
            loader_func = self.supported_extensions.get(extension)
            
            if not loader_func:
                logger.warning(f"No loader for file type: {extension}")
                return None
            
            # Load file content
            content = loader_func(file_path)
            
            if not content or not content.strip():
                logger.warning(f"No content extracted from file: {file_path}")
                return None
            
            # Generate file hash for change detection
            file_hash = self._calculate_file_hash(file_path)
            
            return {
                'id': str(file_path),
                'name': file_path.name,
                'path': str(file_path),
                'content': content,
                'extension': extension,
                'size': stat.st_size,
                'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_hash': file_hash,
                'source': 'local_file',
                'source_id': str(file_path),
                'url': f"file://{file_path.absolute()}"
            }
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file for change detection.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MD5 hash string
        """
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.warning(f"Failed to calculate hash for {file_path}: {e}")
            return ""
    
    def _load_text_file(self, file_path: Path) -> str:
        """Load content from a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            File content as string
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    
    def _load_pdf_file(self, file_path: Path) -> str:
        """Load content from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip()
            
        except ImportError:
            logger.error("PyPDF2 not installed. Cannot process PDF files.")
            return ""
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""
    
    def _load_docx_file(self, file_path: Path) -> str:
        """Load content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx not installed. Cannot process DOCX files.")
            return ""
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""
    
    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get information about a specific file without loading content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File information or None if file doesn't exist
        """
        try:
            path = Path(file_path).expanduser()
            
            if not path.exists() or not path.is_file():
                return None
            
            stat = path.stat()
            
            return {
                'path': str(path),
                'name': path.name,
                'extension': path.suffix.lower(),
                'size': stat.st_size,
                'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_hash': self._calculate_file_hash(path),
                'supported': path.suffix.lower() in self.supported_extensions
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return None
    
    def is_file_changed(self, file_path: str, stored_hash: str) -> bool:
        """Check if a file has changed since last indexing.
        
        Args:
            file_path: Path to the file
            stored_hash: Previously stored file hash
            
        Returns:
            True if file has changed or doesn't exist
        """
        try:
            path = Path(file_path).expanduser()
            
            if not path.exists():
                return True
            
            current_hash = self._calculate_file_hash(path)
            return current_hash != stored_hash
            
        except Exception as e:
            logger.warning(f"Error checking file change for {file_path}: {e}")
            return True  # Assume changed if we can't check